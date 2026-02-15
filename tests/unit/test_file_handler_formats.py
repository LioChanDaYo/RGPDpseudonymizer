"""Unit tests for PDF/DOCX format support in file_handler (Story 5.5).

Tests cover:
- PDF text extraction (Task 5.5.7)
- DOCX text extraction (Task 5.5.8)
- Format-aware read dispatcher (Task 5.5.9)
- CLI command changes for PDF/DOCX (Task 5.5.10)
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

from gdpr_pseudonymizer.exceptions import FileProcessingError
from gdpr_pseudonymizer.utils.file_handler import (
    read_docx,
    read_file,
    read_pdf,
)

# ---------------------------------------------------------------------------
# Fixtures for generating test files programmatically
# ---------------------------------------------------------------------------


def _create_pdf(path: Path, pages: list[str]) -> Path:
    """Create a test PDF with the given page texts."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)
    for text in pages:
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, text)
    pdf.output(str(path))
    return path


def _create_docx(
    path: Path,
    *,
    paragraphs: list[str] | None = None,
    header_text: str | None = None,
    footer_text: str | None = None,
    table_data: list[list[str]] | None = None,
) -> Path:
    """Create a test DOCX with optional paragraphs, header, footer, table."""
    doc = DocxDocument()

    if header_text:
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = header_text

    for text in paragraphs or []:
        doc.add_paragraph(text)

    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text

    if footer_text:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = footer_text

    doc.save(str(path))
    return path


# ===========================================================================
# Task 5.5.7: Unit tests for PDF extraction
# ===========================================================================


class TestReadPdf:
    """Tests for read_pdf() function."""

    def test_single_page_pdf(self, tmp_path: Path) -> None:
        """Test reading a simple single-page PDF."""
        pdf_path = _create_pdf(tmp_path / "single.pdf", ["Bonjour le monde"])
        result = read_pdf(str(pdf_path))
        assert "Bonjour le monde" in result

    def test_multi_page_pdf(self, tmp_path: Path) -> None:
        """Test reading a multi-page PDF with double-newline page separator."""
        pdf_path = _create_pdf(
            tmp_path / "multi.pdf",
            ["Page un contenu", "Page deux contenu", "Page trois contenu"],
        )
        result = read_pdf(str(pdf_path))
        assert "Page un contenu" in result
        assert "Page deux contenu" in result
        assert "Page trois contenu" in result
        # Pages should be joined with double newlines
        assert "\n\n" in result

    def test_empty_pdf_warns_scanned(self, tmp_path: Path) -> None:
        """Test that an empty/scanned PDF triggers a warning log and doesn't error."""
        # Create a PDF with pages but no text (simulates scanned PDF)
        pdf = FPDF()
        pdf.add_page()
        pdf.add_page()
        pdf.output(str(tmp_path / "empty.pdf"))

        with patch("gdpr_pseudonymizer.utils.file_handler.logger") as mock_logger:
            result = read_pdf(str(tmp_path / "empty.pdf"))
            # Should return empty or near-empty string without raising
            assert isinstance(result, str)
            # Verify the scanned PDF warning was actually logged
            mock_logger.warning.assert_called_once()
            call_args = mock_logger.warning.call_args
            assert call_args[0][0] == "scanned_pdf_detected"

    def test_corrupt_pdf_raises_error(self, tmp_path: Path) -> None:
        """Test that a corrupt PDF raises FileProcessingError."""
        corrupt_pdf = tmp_path / "corrupt.pdf"
        corrupt_pdf.write_bytes(b"NOT A VALID PDF FILE CONTENT")

        with pytest.raises(FileProcessingError, match="Cannot read PDF file"):
            read_pdf(str(corrupt_pdf))

    def test_password_protected_pdf_raises_error(self, tmp_path: Path) -> None:
        """Test that a password-protected PDF raises FileProcessingError."""
        # Mock pdfplumber to raise an exception with 'password' in the message
        pdf_path = _create_pdf(tmp_path / "protected.pdf", ["test"])

        with patch(
            "gdpr_pseudonymizer.utils.file_handler.pdfplumber"
        ) as mock_pdfplumber:
            mock_pdfplumber.open.side_effect = Exception("PDF is password protected")
            with pytest.raises(FileProcessingError, match="password-protected"):
                read_pdf(str(pdf_path))

    def test_pdfplumber_not_installed(self) -> None:
        """Test error message when pdfplumber is not installed."""
        with patch("gdpr_pseudonymizer.utils.file_handler.HAS_PDFPLUMBER", False):
            with pytest.raises(
                FileProcessingError,
                match="PDF support requires 'pdfplumber'",
            ):
                read_pdf("any_file.pdf")


# ===========================================================================
# Task 5.5.8: Unit tests for DOCX extraction
# ===========================================================================


class TestReadDocx:
    """Tests for read_docx() function."""

    def test_simple_paragraphs(self, tmp_path: Path) -> None:
        """Test extracting paragraphs from a simple DOCX."""
        docx_path = _create_docx(
            tmp_path / "simple.docx",
            paragraphs=["Premier paragraphe.", "Deuxieme paragraphe."],
        )
        result = read_docx(str(docx_path))
        assert "Premier paragraphe." in result
        assert "Deuxieme paragraphe." in result

    def test_docx_with_header(self, tmp_path: Path) -> None:
        """Test extracting header text from DOCX."""
        docx_path = _create_docx(
            tmp_path / "header.docx",
            paragraphs=["Body text"],
            header_text="En-tete du document",
        )
        result = read_docx(str(docx_path))
        assert "En-tete du document" in result
        assert "Body text" in result

    def test_docx_with_footer(self, tmp_path: Path) -> None:
        """Test extracting footer text from DOCX."""
        docx_path = _create_docx(
            tmp_path / "footer.docx",
            paragraphs=["Body text"],
            footer_text="Pied de page",
        )
        result = read_docx(str(docx_path))
        assert "Body text" in result
        assert "Pied de page" in result

    def test_docx_with_table(self, tmp_path: Path) -> None:
        """Test extracting table text from DOCX."""
        docx_path = _create_docx(
            tmp_path / "table.docx",
            paragraphs=["Avant le tableau"],
            table_data=[
                ["Nom", "Ville"],
                ["Jean Dupont", "Paris"],
                ["Marie Martin", "Lyon"],
            ],
        )
        result = read_docx(str(docx_path))
        assert "Avant le tableau" in result
        assert "Jean Dupont" in result
        assert "Paris" in result
        assert "Marie Martin" in result
        assert "Lyon" in result

    def test_composite_docx_all_elements(self, tmp_path: Path) -> None:
        """Test extracting all element types from a single DOCX (header+paragraphs+table+footer)."""
        docx_path = _create_docx(
            tmp_path / "composite.docx",
            paragraphs=["Premier paragraphe.", "Deuxieme paragraphe."],
            header_text="En-tete confidentiel",
            footer_text="Page 1 sur 1",
            table_data=[
                ["Nom", "Prenom"],
                ["Dupont", "Jean"],
            ],
        )
        result = read_docx(str(docx_path))
        # Header
        assert "En-tete confidentiel" in result
        # Body paragraphs
        assert "Premier paragraphe." in result
        assert "Deuxieme paragraphe." in result
        # Table content
        assert "Dupont" in result
        assert "Jean" in result
        # Footer
        assert "Page 1 sur 1" in result

    def test_corrupt_docx_raises_error(self, tmp_path: Path) -> None:
        """Test that a corrupt DOCX raises FileProcessingError."""
        corrupt_docx = tmp_path / "corrupt.docx"
        corrupt_docx.write_bytes(b"NOT A VALID DOCX FILE")

        with pytest.raises(FileProcessingError, match="Cannot read DOCX file"):
            read_docx(str(corrupt_docx))

    def test_python_docx_not_installed(self) -> None:
        """Test error message when python-docx is not installed."""
        with patch("gdpr_pseudonymizer.utils.file_handler.HAS_PYTHON_DOCX", False):
            with pytest.raises(
                FileProcessingError,
                match="DOCX support requires 'python-docx'",
            ):
                read_docx("any_file.docx")


# ===========================================================================
# Task 5.5.9: Unit tests for format detection and dispatcher
# ===========================================================================


class TestReadFileDispatcher:
    """Tests for read_file() format-aware dispatch."""

    def test_dispatches_pdf(self, tmp_path: Path) -> None:
        """Test read_file dispatches .pdf to read_pdf."""
        pdf_path = _create_pdf(tmp_path / "test.pdf", ["PDF content here"])
        result = read_file(str(pdf_path))
        assert "PDF content here" in result

    def test_dispatches_docx(self, tmp_path: Path) -> None:
        """Test read_file dispatches .docx to read_docx."""
        docx_path = _create_docx(
            tmp_path / "test.docx", paragraphs=["DOCX content here"]
        )
        result = read_file(str(docx_path))
        assert "DOCX content here" in result

    def test_txt_still_works(self, tmp_path: Path) -> None:
        """Test read_file still works for .txt (regression)."""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Plain text content", encoding="utf-8")
        result = read_file(str(txt_file))
        assert result == "Plain text content"

    def test_md_still_works(self, tmp_path: Path) -> None:
        """Test read_file still works for .md (regression)."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Markdown heading", encoding="utf-8")
        result = read_file(str(md_file))
        assert result == "# Markdown heading"

    def test_case_insensitive_pdf(self, tmp_path: Path) -> None:
        """Test case-insensitive .PDF extension handling."""
        pdf_path = _create_pdf(tmp_path / "test.PDF", ["Upper case extension"])
        result = read_file(str(pdf_path))
        assert "Upper case extension" in result

    def test_case_insensitive_docx(self, tmp_path: Path) -> None:
        """Test case-insensitive .Docx extension handling."""
        docx_path = _create_docx(
            tmp_path / "test.Docx", paragraphs=["Mixed case extension"]
        )
        result = read_file(str(docx_path))
        assert "Mixed case extension" in result

    def test_file_not_found(self, tmp_path: Path) -> None:
        """Test read_file raises error for non-existent file."""
        with pytest.raises(FileProcessingError, match="File not found"):
            read_file(str(tmp_path / "nonexistent.pdf"))

    def test_path_is_directory(self, tmp_path: Path) -> None:
        """Test read_file raises error when path is a directory."""
        with pytest.raises(FileProcessingError, match="Path is not a file"):
            read_file(str(tmp_path))


# ===========================================================================
# Task 5.5.10: Unit tests for CLI command changes
# ===========================================================================


class TestProcessCommandFormats:
    """Tests for process command PDF/DOCX acceptance."""

    def test_process_accepts_pdf_extension(self, tmp_path: Path) -> None:
        """Test process command accepts .pdf files."""
        from unittest.mock import MagicMock

        import typer
        from typer.testing import CliRunner

        from gdpr_pseudonymizer.cli.commands.process import process_command

        test_app = typer.Typer()

        @test_app.callback()
        def _cb() -> None:
            pass

        test_app.command(name="process")(process_command)
        cli_runner = CliRunner()

        # Create a real PDF file
        pdf_path = _create_pdf(tmp_path / "input.pdf", ["Test content"])

        with (
            patch(
                "gdpr_pseudonymizer.cli.commands.process.resolve_passphrase"
            ) as mock_pp,
            patch(
                "gdpr_pseudonymizer.cli.commands.process.DocumentProcessor"
            ) as mock_dp,
            patch("gdpr_pseudonymizer.cli.validators.init_database"),
        ):
            mock_pp.return_value = "testpassphrase123!"
            mock_result = MagicMock(
                success=True,
                entities_detected=0,
                entities_new=0,
                entities_reused=0,
                processing_time_seconds=0.1,
            )
            mock_dp.return_value.process_document.return_value = mock_result

            result = cli_runner.invoke(test_app, ["process", str(pdf_path)])

        assert result.exit_code == 0

    def test_process_accepts_docx_extension(self, tmp_path: Path) -> None:
        """Test process command accepts .docx files."""
        from unittest.mock import MagicMock

        import typer
        from typer.testing import CliRunner

        from gdpr_pseudonymizer.cli.commands.process import process_command

        test_app = typer.Typer()

        @test_app.callback()
        def _cb() -> None:
            pass

        test_app.command(name="process")(process_command)
        cli_runner = CliRunner()

        # Create a real DOCX file
        docx_path = _create_docx(tmp_path / "input.docx", paragraphs=["Test content"])

        with (
            patch(
                "gdpr_pseudonymizer.cli.commands.process.resolve_passphrase"
            ) as mock_pp,
            patch(
                "gdpr_pseudonymizer.cli.commands.process.DocumentProcessor"
            ) as mock_dp,
            patch("gdpr_pseudonymizer.cli.validators.init_database"),
        ):
            mock_pp.return_value = "testpassphrase123!"
            mock_result = MagicMock(
                success=True,
                entities_detected=0,
                entities_new=0,
                entities_reused=0,
                processing_time_seconds=0.1,
            )
            mock_dp.return_value.process_document.return_value = mock_result

            result = cli_runner.invoke(test_app, ["process", str(docx_path)])

        assert result.exit_code == 0

    def test_default_output_txt_for_pdf_input(self, tmp_path: Path) -> None:
        """Test that default output filename uses .txt for PDF input."""
        from unittest.mock import MagicMock

        import typer
        from typer.testing import CliRunner

        from gdpr_pseudonymizer.cli.commands.process import process_command

        test_app = typer.Typer()

        @test_app.callback()
        def _cb() -> None:
            pass

        test_app.command(name="process")(process_command)
        cli_runner = CliRunner()

        pdf_path = _create_pdf(tmp_path / "report.pdf", ["Content"])

        with (
            patch(
                "gdpr_pseudonymizer.cli.commands.process.resolve_passphrase"
            ) as mock_pp,
            patch(
                "gdpr_pseudonymizer.cli.commands.process.DocumentProcessor"
            ) as mock_dp,
            patch("gdpr_pseudonymizer.cli.validators.init_database"),
        ):
            mock_pp.return_value = "testpassphrase123!"
            mock_result = MagicMock(
                success=True,
                entities_detected=0,
                entities_new=0,
                entities_reused=0,
                processing_time_seconds=0.1,
            )
            mock_dp.return_value.process_document.return_value = mock_result

            result = cli_runner.invoke(test_app, ["process", str(pdf_path)])

        assert result.exit_code == 0
        # Check that process_document was called with .txt output
        call_args = mock_dp.return_value.process_document.call_args
        output_path = call_args[1].get("output_path") or call_args[0][1]
        assert output_path.endswith("_pseudonymized.txt")

    def test_default_output_txt_for_docx_input(self, tmp_path: Path) -> None:
        """Test that default output filename uses .txt for DOCX input."""
        from unittest.mock import MagicMock

        import typer
        from typer.testing import CliRunner

        from gdpr_pseudonymizer.cli.commands.process import process_command

        test_app = typer.Typer()

        @test_app.callback()
        def _cb() -> None:
            pass

        test_app.command(name="process")(process_command)
        cli_runner = CliRunner()

        docx_path = _create_docx(tmp_path / "report.docx", paragraphs=["Content"])

        with (
            patch(
                "gdpr_pseudonymizer.cli.commands.process.resolve_passphrase"
            ) as mock_pp,
            patch(
                "gdpr_pseudonymizer.cli.commands.process.DocumentProcessor"
            ) as mock_dp,
            patch("gdpr_pseudonymizer.cli.validators.init_database"),
        ):
            mock_pp.return_value = "testpassphrase123!"
            mock_result = MagicMock(
                success=True,
                entities_detected=0,
                entities_new=0,
                entities_reused=0,
                processing_time_seconds=0.1,
            )
            mock_dp.return_value.process_document.return_value = mock_result

            result = cli_runner.invoke(test_app, ["process", str(docx_path)])

        assert result.exit_code == 0
        call_args = mock_dp.return_value.process_document.call_args
        output_path = call_args[1].get("output_path") or call_args[0][1]
        assert output_path.endswith("_pseudonymized.txt")

    def test_unsupported_extension_rejected(self, tmp_path: Path) -> None:
        """Test that unsupported extensions are still rejected."""
        csv_file = tmp_path / "data.csv"
        csv_file.write_text("a,b,c")

        import typer
        from typer.testing import CliRunner

        from gdpr_pseudonymizer.cli.commands.process import process_command

        test_app = typer.Typer()

        @test_app.callback()
        def _cb() -> None:
            pass

        test_app.command(name="process")(process_command)
        cli_runner = CliRunner()

        result = cli_runner.invoke(test_app, ["process", str(csv_file)])
        # Typer validates exists=True, so for an existing csv it should reach our check
        assert result.exit_code == 1


class TestBatchCommandFormats:
    """Tests for batch command PDF/DOCX collection."""

    def test_collect_files_includes_pdf(self, tmp_path: Path) -> None:
        """Test collect_files includes .pdf files."""
        from gdpr_pseudonymizer.cli.commands.batch import collect_files

        _create_pdf(tmp_path / "test.pdf", ["content"])
        files = collect_files(tmp_path)
        assert any(f.name == "test.pdf" for f in files)

    def test_collect_files_includes_docx(self, tmp_path: Path) -> None:
        """Test collect_files includes .docx files."""
        from gdpr_pseudonymizer.cli.commands.batch import collect_files

        _create_docx(tmp_path / "test.docx", paragraphs=["content"])
        files = collect_files(tmp_path)
        assert any(f.name == "test.docx" for f in files)

    def test_collect_files_mixed_formats(self, tmp_path: Path) -> None:
        """Test collect_files handles mixed format directories."""
        from gdpr_pseudonymizer.cli.commands.batch import collect_files

        (tmp_path / "a.txt").write_text("txt")
        (tmp_path / "b.md").write_text("md")
        _create_pdf(tmp_path / "c.pdf", ["pdf"])
        _create_docx(tmp_path / "d.docx", paragraphs=["docx"])
        (tmp_path / "e.csv").write_text("csv")  # Unsupported

        files = collect_files(tmp_path)
        names = {f.name for f in files}
        assert names == {"a.txt", "b.md", "c.pdf", "d.docx"}

    def test_dependency_not_installed_error_message(self) -> None:
        """Test error message when PDF dependency is missing."""
        with patch("gdpr_pseudonymizer.utils.file_handler.HAS_PDFPLUMBER", False):
            with pytest.raises(
                FileProcessingError,
                match="pip install gdpr-pseudonymizer\\[pdf\\]",
            ):
                read_pdf("test.pdf")

    def test_docx_dependency_not_installed_error_message(self) -> None:
        """Test error message when DOCX dependency is missing."""
        with patch("gdpr_pseudonymizer.utils.file_handler.HAS_PYTHON_DOCX", False):
            with pytest.raises(
                FileProcessingError,
                match="pip install gdpr-pseudonymizer\\[docx\\]",
            ):
                read_docx("test.docx")
