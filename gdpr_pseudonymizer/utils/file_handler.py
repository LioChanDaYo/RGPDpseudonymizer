"""File handling utilities with cross-platform path support.

This module provides safe file I/O operations using pathlib
for cross-platform compatibility. Supports .txt, .md, .pdf, .docx, .xlsx, and .csv formats.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path

from gdpr_pseudonymizer.exceptions import FileProcessingError
from gdpr_pseudonymizer.utils.logger import get_logger

logger = get_logger(__name__)

# Optional dependency flags
try:
    import pdfplumber

    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import docx  # python-docx

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import openpyxl  # type: ignore[import-untyped]

    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

# Minimum characters per page to consider a PDF as having text content
_SCANNED_PDF_CHARS_PER_PAGE = 50


def read_file(file_path: str) -> str:
    """Read file contents, auto-detecting format from extension.

    Supports .txt, .md (plaintext), .pdf (via pdfplumber), and .docx (via python-docx).

    Args:
        file_path: Path to file (absolute or relative)

    Returns:
        File contents as string

    Raises:
        FileProcessingError: If file not found, cannot be read, or
            required optional dependency is missing
    """
    path = Path(file_path)

    if not path.exists():
        raise FileProcessingError(f"File not found: {file_path}")

    if not path.is_file():
        raise FileProcessingError(f"Path is not a file: {file_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return read_pdf(file_path)

    if ext == ".docx":
        return read_docx(file_path)

    if ext == ".xlsx":
        return read_excel(file_path)

    if ext == ".csv":
        return read_csv(file_path)

    # Default: plaintext (.txt, .md, or any other text file)
    try:
        with open(path, encoding="utf-8") as f:
            return f.read()
    except PermissionError as e:
        raise FileProcessingError(f"Permission denied: {file_path}") from e
    except OSError as e:
        raise FileProcessingError(f"Cannot read file {file_path}: {e}") from e


def read_pdf(file_path: str) -> str:
    """Extract text content from a PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text with pages joined by double newlines

    Raises:
        FileProcessingError: If pdfplumber not installed, file corrupt,
            or file is password-protected
    """
    if not HAS_PDFPLUMBER:
        raise FileProcessingError(
            "PDF support requires 'pdfplumber'. "
            "Install with: pip install gdpr-pseudonymizer[pdf]"
        )

    path = Path(file_path)

    try:
        pages_text: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)

            page_count = len(pdf.pages)

        extracted = "\n\n".join(pages_text)

        # Detect scanned/image-based PDFs
        if (
            page_count > 0
            and len(extracted.strip()) < _SCANNED_PDF_CHARS_PER_PAGE * page_count
        ):
            logger.warning(
                "scanned_pdf_detected",
                file=file_path,
                page_count=page_count,
                extracted_chars=len(extracted.strip()),
                message=(
                    "This PDF appears to be scanned/image-based. "
                    "Text extraction may be incomplete."
                ),
            )

        return extracted

    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            raise FileProcessingError(
                f"PDF is password-protected: {file_path}. "
                "Please provide an unprotected PDF."
            ) from e
        raise FileProcessingError(f"Cannot read PDF file {file_path}: {e}") from e


def read_docx(file_path: str) -> str:
    """Extract text content from a DOCX file.

    Extracts text from paragraphs, headers, footers, and tables (best-effort).

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text with sections joined by newlines

    Raises:
        FileProcessingError: If python-docx not installed or file corrupt
    """
    if not HAS_PYTHON_DOCX:
        raise FileProcessingError(
            "DOCX support requires 'python-docx'. "
            "Install with: pip install gdpr-pseudonymizer[docx]"
        )

    path = Path(file_path)

    try:
        document = docx.Document(str(path))
        text_parts: list[str] = []

        # Extract header text (best-effort)
        for section in document.sections:
            header = section.header
            if header and not header.is_linked_to_previous:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

        # Extract body paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract table text (best-effort: row-by-row, cell text joined)
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        # Extract footer text (best-effort)
        for section in document.sections:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

        return "\n".join(text_parts)

    except Exception as e:
        raise FileProcessingError(f"Cannot read DOCX file {file_path}: {e}") from e


def read_excel(file_path: str) -> str:
    """Extract text content from an Excel (.xlsx) file.

    Args:
        file_path: Path to Excel file

    Returns:
        Extracted text with cells tab-separated, rows newline-separated,
        sheets double-newline-separated

    Raises:
        FileProcessingError: If openpyxl not installed, file corrupt,
            password-protected, or binary .xls format
    """
    if not HAS_OPENPYXL:
        raise FileProcessingError(
            "Excel support requires 'openpyxl'. "
            "Install with: pip install gdpr-pseudonymizer[excel]"
        )

    path = Path(file_path)

    # Check for unsupported binary .xls format
    if path.suffix.lower() == ".xls":
        raise FileProcessingError(
            "Binary .xls format is not supported. Please convert to .xlsx."
        )

    try:
        wb = openpyxl.load_workbook(str(path), data_only=True)
    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            raise FileProcessingError(
                f"Excel file is password-protected: {file_path}. "
                "Please provide an unprotected file."
            ) from e
        if "invalid" in error_msg or "not a zip" in error_msg or "corrupt" in error_msg:
            raise FileProcessingError(
                f"Cannot read Excel file {file_path}: "
                "File appears to be corrupt or not a valid .xlsx file. "
                "Binary .xls format is not supported. Please convert to .xlsx."
            ) from e
        raise FileProcessingError(f"Cannot read Excel file {file_path}: {e}") from e

    try:
        # Count non-empty cells for large-file warning
        non_empty_count = 0
        sheets_text: list[str] = []

        for sheet in wb.worksheets:
            rows_text: list[str] = []
            for row in sheet.iter_rows():
                cells: list[str] = []
                for cell in row:
                    if cell.value is not None:
                        cells.append(str(cell.value))
                        non_empty_count += 1
                if cells:
                    rows_text.append("\t".join(cells))
            if rows_text:
                sheets_text.append("\n".join(rows_text))

        if non_empty_count > 100_000:
            logger.warning(
                "large_excel_file",
                file=file_path,
                non_empty_cells=non_empty_count,
                message=(
                    "This Excel file contains more than 100,000 non-empty cells. "
                    "Processing may consume significant memory."
                ),
            )

        return "\n\n".join(sheets_text)
    finally:
        wb.close()


def read_csv(file_path: str) -> str:
    """Extract text content from a CSV file.

    Auto-detects encoding (UTF-8, then Latin-1 fallback) and delimiter.

    Args:
        file_path: Path to CSV file

    Returns:
        Extracted text with cells tab-separated, rows newline-separated

    Raises:
        FileProcessingError: If file cannot be read
    """
    path = Path(file_path)

    # Auto-detect encoding: try UTF-8 first, fall back to latin-1
    content: str | None = None
    try:
        raw_bytes = path.read_bytes()
    except PermissionError as e:
        raise FileProcessingError(f"Permission denied: {file_path}") from e
    except OSError as e:
        raise FileProcessingError(f"Cannot read file {file_path}: {e}") from e

    try:
        content = raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        content = raw_bytes.decode("latin-1")

    if not content.strip():
        return ""

    # Auto-detect delimiter using csv.Sniffer on first 8KB
    sample = content[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample)
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = ","

    reader = csv.reader(io.StringIO(content), delimiter=delimiter)
    rows_text: list[str] = []
    for row in reader:
        if any(cell.strip() for cell in row):
            rows_text.append("\t".join(row))

    return "\n".join(rows_text)


def write_file(file_path: str, content: str) -> None:
    """Write text content to file.

    Args:
        file_path: Path to file (absolute or relative)
        content: Text content to write

    Raises:
        FileProcessingError: If file cannot be written
    """
    path = Path(file_path)

    try:
        # Create parent directories if they don't exist
        path.parent.mkdir(parents=True, exist_ok=True)

        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    except PermissionError as e:
        raise FileProcessingError(f"Permission denied: {file_path}") from e
    except OSError as e:
        raise FileProcessingError(f"Cannot write file {file_path}: {e}") from e


def validate_file_path(file_path: str, allowed_extensions: list[str]) -> bool:
    """Validate file path and extension.

    Args:
        file_path: Path to validate
        allowed_extensions: List of allowed file extensions (e.g., ['.txt', '.md'])

    Returns:
        True if path is valid and has allowed extension

    Raises:
        FileProcessingError: If path is invalid or extension not allowed
    """
    path = Path(file_path)

    # Check if path is absolute or can be resolved
    if not path.is_absolute():
        path = path.resolve()

    # Check file extension
    if path.suffix.lower() not in allowed_extensions:
        raise FileProcessingError(
            f"Invalid file extension: {path.suffix}. "
            f"Allowed: {', '.join(allowed_extensions)}"
        )

    return True


def get_file_extension(file_path: str) -> str:
    """Get file extension from path.

    Args:
        file_path: Path to file

    Returns:
        File extension with dot (e.g., '.txt', '.md')
    """
    return Path(file_path).suffix.lower()


def ensure_absolute_path(file_path: str) -> str:
    """Convert relative path to absolute path.

    Args:
        file_path: Path to convert

    Returns:
        Absolute path as string
    """
    path = Path(file_path)
    # resolve() always returns absolute path, even for non-existent files
    return str(path.resolve())
